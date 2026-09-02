
import traceback
import os, uuid, asyncio, requests, re, json, traceback
from typing import Dict, List
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv, dotenv_values
from datetime import datetime
from pathlib import Path

load_dotenv()
app = FastAPI(title="Council AI V20 FINAL V4 - Bien ban chuan VN + Word/TXT")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)
PERSONNEL_FILE = DATA_DIR / "personnel.json"
MEETINGS_DIR = DATA_DIR / "meetings"
MEETINGS_DIR.mkdir(exist_ok=True)
FRONTEND_FILE = Path(__file__).parent / "council-v20-frontend.html"

councils_db: Dict = {}
ALIVE_PROVIDERS: List[Dict] = []

def scan_env_apis():
    env = {}
    if os.path.exists(".env"):
        env.update(dotenv_values(".env"))
    env.update(dict(os.environ))
    apis_raw = []
    for k, v in env.items():
        if not v or len(v.strip()) < 10: continue
        uk = k.upper()
        if "API_KEY" not in uk: continue
        if uk in ["ANDROID_API_KEY"]: continue
        provider = uk.replace("_API_KEY","").lower()
        if provider in ["android"]: continue
        prov_base = provider.split("_")[0]
        if "mistral" in provider: prov_base = "mistral"
        elif "groq" in provider: prov_base = "groq"
        elif "openrouter" in provider: prov_base = "openrouter"
        else:
            if prov_base not in ["mistral","groq","openrouter"]: continue
        base_url = env.get(f"{provider.upper()}_API_URL") or env.get(f"{prov_base.upper()}_API_URL") or ""
        model = env.get(f"{provider.upper()}_MODEL") or env.get(f"{prov_base.upper()}_MODEL") or ""
        if not base_url:
            if "mistral" in prov_base: base_url = "https://api.mistral.ai/v1"
            elif "groq" in prov_base: base_url = "https://api.groq.com/openai/v1"
            elif "openrouter" in prov_base: base_url = "https://openrouter.ai/api/v1"
        apis_raw.append({"provider": prov_base, "api_key": v.strip(), "base_url": base_url.strip().rstrip("/"), "model": model.strip(), "env_key": k})
    deduped = {}
    for api in apis_raw:
        pb = api["provider"]
        if pb not in deduped:
            deduped[pb] = api
        else:
            if len(api["model"]) > len(deduped[pb]["model"]):
                deduped[pb] = api
    return list(deduped.values())

def test_provider_alive(api: Dict) -> bool:
    p, key, base, model = api["provider"], api["api_key"], api["base_url"], api["model"]
    try:
        if "mistral" in p:
            m = model or "mistral-small-latest"
            r = requests.post(f"{base or 'https://api.mistral.ai/v1'}/chat/completions", headers={"Authorization": f"Bearer {key}", "Content-Type":"application/json"}, json={"model": m, "messages": [{"role":"user","content":"hi"}], "max_tokens":5}, timeout=12)
            if r.status_code in [200,429]:
                api["model"]=m; api["status"]="alive" if r.status_code==200 else "rate_limited"; return True
            return False
        elif "groq" in p:
            for m in [model, "openai/gpt-oss-120b", "llama-3.3-70b-versatile"]:
                if not m: continue
                r = requests.post("https://api.groq.com/openai/v1/chat/completions", headers={"Authorization": f"Bearer {key}", "Content-Type":"application/json"}, json={"model": m, "messages": [{"role":"user","content":"hi"}], "max_tokens":5}, timeout=12)
                if r.status_code in [200,429]:
                    api["model"]=m; api["status"]="alive" if r.status_code==200 else "rate_limited"; return True
            return False
        elif "openrouter" in p:
            headers = {"Authorization": f"Bearer {key}", "Content-Type":"application/json", "HTTP-Referer":"http://localhost:3000", "X-Title":"Council AI"}
            for m in [model, "openai/gpt-latest", "meta-llama/llama-3.3-70b-instruct:free"]:
                if not m: continue
                r = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json={"model": m, "messages":[{"role":"user","content":"hi"}], "max_tokens":5}, timeout=15)
                if r.status_code in [200,429]:
                    api["model"]=m; api["status"]="alive" if r.status_code==200 else "rate_limited"; return True
            return False
        else: return False
    except: return False

def refresh_alive_providers():
    global ALIVE_PROVIDERS
    all_apis = scan_env_apis()
    alive = []
    for api in all_apis:
        if test_provider_alive(api):
            alive.append(api)
    if not alive and all_apis:
        alive = [all_apis[0]]; alive[0]["status"]="fallback"
    ALIVE_PROVIDERS = alive
    print(f"\n🔥 {len(ALIVE_PROVIDERS)} API SONG THAT: {[p['provider'] for p in ALIVE_PROVIDERS]}")
    return ALIVE_PROVIDERS

refresh_alive_providers()

def get_default_personnel():
    prov = lambda i: ALIVE_PROVIDERS[i % len(ALIVE_PROVIDERS)]["provider"] if ALIVE_PROVIDERS else "mistral"
    mod = lambda i: ALIVE_PROVIDERS[i % len(ALIVE_PROVIDERS)]["model"] if ALIVE_PROVIDERS else "mistral-small-latest"
    return [
        {"id": "core_thuky", "name": "Thư Ký Hội Đồng", "role": "Thư Ký Hội Đồng", "desc": "Ghi biên bản, điểm danh, lưu file", "system_prompt": "Bạn là Thư Ký Hội Đồng - vai CORE. Điểm danh, ghi biên bản 80-120 từ, lưu file.", "color": "#2B2D42", "type": "core", "core_type": "thuky", "provider": prov(0), "model": mod(0), "trained": True},
        {"id": "core_phoql", "name": "Phó Quản Lý Điều Hành", "role": "Phó Quản Lý Điều Hành", "desc": "Tổng hợp, ra quyết định cuối, 3 bước + deadline", "system_prompt": "Bạn là Phó Quản Lý Điều Hành - vai CORE quyền lực thứ 2. Tổng hợp ý kiến, ra quyết định cuối cùng, 3 bước hành động có deadline.", "color": "#EF476F", "type": "core", "core_type": "phoql", "provider": prov(1) if len(ALIVE_PROVIDERS)>1 else prov(0), "model": mod(1) if len(ALIVE_PROVIDERS)>1 else mod(0), "trained": True},
        {"id": "core_nhansu", "name": "Trưởng Phòng Nhân Sự", "role": "Trưởng Phòng Nhân Sự", "desc": "Quản lý Danh sách Hội đồng, đặt chức danh và prompt cho vai", "system_prompt": "Bạn là Trưởng Phòng Nhân Sự - vai CORE quản lý Danh sách Hội đồng. Nhiệm vụ: đặt chức danh theo nhu cầu, viết system_prompt riêng cho từng chức danh, đề xuất vai phù hợp ngữ cảnh.", "color": "#118AB2", "type": "core", "core_type": "nhansu", "provider": prov(2) if len(ALIVE_PROVIDERS)>2 else prov(0), "model": mod(2) if len(ALIVE_PROVIDERS)>2 else mod(0), "trained": True},
    ]

ROLE_TEMPLATES = [
    {"role": "Chuyên gia Kỹ Thuật", "desc": "Giải pháp kỹ thuật", "prompt": "Bạn là Chuyên gia Kỹ Thuật theo chức danh. Đưa giải pháp kỹ thuật cụ thể, chi tiết 200-300 từ.", "color": "#6EC1E4"},
    {"role": "Chuyên gia Marketing", "desc": "Chiến lược viral", "prompt": "Bạn là Chuyên gia Marketing theo chức danh. Đưa chiến lược viral chi tiết 200-300 từ.", "color": "#FF8FA3"},
    {"role": "Chuyên gia Dữ Liệu", "desc": "Phân tích số liệu", "prompt": "Bạn là Chuyên gia Dữ Liệu theo chức danh. Đưa số liệu, insight chi tiết 200-300 từ.", "color": "#A78BFA"},
    {"role": "Chuyên gia Tài Chính", "desc": "Tối ưu chi phí", "prompt": "Bạn là Chuyên gia Tài Chính theo chức danh. Phân tích chi phí, ROI chi tiết 200-300 từ.", "color": "#06D6A0"},
]

def load_personnel():
    if PERSONNEL_FILE.exists():
        try:
            content = PERSONNEL_FILE.read_text(encoding='utf-8').strip()
            if not content:
                raise ValueError("personnel.json empty")
            data = json.loads(content)
            seen={}; cleaned=[]
            # Sap xep CORE truoc, giu nguyen thu tu tao cho chuyen gia tu dat
            data_sorted=sorted(data, key=lambda x: (0 if x.get("type")=="core" else 1, x.get("created_at","")))
            for p in data_sorted:
                rk=p.get("role","").strip().lower()
                if not rk: continue
                # Chuyen gia tu dat co the co ky tu dac biet &, khong bo
                if rk in seen: 
                    # Neu trung ten thi giu ban moi nhat co system_prompt dai hon
                    if len(p.get("system_prompt","")) > len(seen[rk].get("system_prompt","")):
                        # Thay the
                        for idx, cp in enumerate(cleaned):
                            if cp.get("role","").strip().lower()==rk:
                                cleaned[idx]=p
                                break
                        seen[rk]=p
                    continue
                dup=False
                for ek in seen:
                    # Chi dedup CORE thoi, chuyen gia tu dat khong dedup
                    if p.get("type")=="core" and seen[ek].get("type")=="core":
                        if ("thu ky" in rk and "thu ky" in ek) or ("phó quản lý" in rk and "phó quản lý" in ek) or ("pho quan ly" in rk and "pho quan ly" in ek) or ("nhân sự" in rk and "nhân sự" in ek) or ("nhan su" in rk and "nhan su" in ek):
                            dup=True; break
                if dup: continue
                seen[rk]=p; cleaned.append(p)
            defaults=get_default_personnel()
            core_present={p.get("core_type") for p in cleaned if p.get("type")=="core"}
            final=[p for p in cleaned if p["id"] not in ["p_thuky","p_phoql"]]
            for d in defaults:
                if d["core_type"] not in core_present:
                    final.insert(0,d)
                else:
                    for i,ex in enumerate(final):
                        if ex.get("core_type")==d["core_type"]:
                            # Giu lai system_prompt tu dat neu co?
                            if len(ex.get("system_prompt",""))>50 and ex.get("auto_generated")!=True:
                                # Giua nguyen nhung cap nhat provider
                                ex["provider"]=d["provider"]
                                ex["model"]=d["model"]
                            else:
                                final[i]=d
                            break

            # === FIX: Phan deu 3 API cho cac chuyen gia, khong de chi dung mistral ===
            # Neu tat ca chuyen gia deu dung 1 provider (nhu anh Boss chup), chia deu lai
            council_roles = [p for p in final if p.get("type")!="core"]
            if council_roles:
                providers_used = set([p.get("provider","") for p in council_roles])
                if len(providers_used) <= 1 and len(ALIVE_PROVIDERS)>=2:
                    # Chia deu lai theo vong tron
                    for idx, p in enumerate(council_roles):
                        prov_idx = idx % len(ALIVE_PROVIDERS)
                        p["provider"] = ALIVE_PROVIDERS[prov_idx]["provider"]
                        p["model"] = ALIVE_PROVIDERS[prov_idx]["model"]
                    print(f"🔄 Da phan deu lai {len(council_roles)} chuyen gia cho {len(ALIVE_PROVIDERS)} API: {[p['provider'] for p in council_roles]}")
                    # Cap nhat final
                    # final da chua tham chieu den council_roles nen ok

            PERSONNEL_FILE.write_text(json.dumps(final, ensure_ascii=False, indent=2), encoding='utf-8')
            return final
        except Exception as e:
            print(f"Reset: {e}")
            traceback.print_exc()
            # Return default if reset fails
            default=get_default_personnel()
            Path(PERSONNEL_FILE).write_text(json.dumps(default, ensure_ascii=False, indent=2), encoding='utf-8')
            return default
    default=get_default_personnel()
    PERSONNEL_FILE.write_text(json.dumps(default, ensure_ascii=False, indent=2), encoding='utf-8')
    return default

def get_provider_info(name: str):
    for p in ALIVE_PROVIDERS:
        if p["provider"]==name.lower():
            return p
    return ALIVE_PROVIDERS[0] if ALIVE_PROVIDERS else {"provider":"mistral","api_key":"demo","base_url":"https://api.mistral.ai/v1","model":"mistral-small-latest","status":"demo"}

def call_llm_real(provider_info, system_prompt, user_prompt, max_tokens=600, temp=0.7):
    try:
        if not provider_info or provider_info.get("api_key")=="demo":
            return f"[{provider_info.get('provider','demo')}] Fallback cho: {user_prompt[:120]}"
        base=provider_info.get("base_url") or "https://api.mistral.ai/v1"
        model=provider_info.get("model") or "mistral-small-latest"
        key=provider_info.get("api_key") or ""
        msgs=[{"role":"system","content":system_prompt},{"role":"user","content":user_prompt}]
        headers={"Authorization": f"Bearer {key}", "Content-Type":"application/json"}
        if "openrouter" in provider_info.get("provider",""):
            headers.update({"HTTP-Referer":"http://localhost:3000","X-Title":"Council AI"})
        r=requests.post(f"{base}/chat/completions", headers=headers, json={"model":model,"messages":msgs,"max_tokens":max_tokens,"temperature":temp}, timeout=30)
        if r.status_code==200 and "choices" in r.json():
            return r.json()["choices"][0]["message"]["content"]
        for alt in ALIVE_PROVIDERS:
            if alt["provider"]!=provider_info["provider"]:
                try:
                    b2=alt.get("base_url") or base; m2=alt.get("model") or model; k2=alt.get("api_key")
                    h2={"Authorization": f"Bearer {k2}", "Content-Type":"application/json"}
                    if "openrouter" in alt.get("provider",""): h2.update({"HTTP-Referer":"http://localhost:3000","X-Title":"Council AI"})
                    r2=requests.post(f"{b2}/chat/completions", headers=h2, json={"model":m2,"messages":msgs,"max_tokens":max_tokens,"temperature":temp}, timeout=25)
                    if r2.status_code==200 and "choices" in r2.json():
                        return r2.json()["choices"][0]["message"]["content"] + f" (via {alt['provider']})"
                except: continue
        return f"[FALLBACK] {user_prompt[:120]}"
    except Exception as e:
        return f"[FALLBACK EX] {str(e)[:80]}"

class CreateCouncilReq(BaseModel):
    problem: str
    selected_personnel_ids: List[str]
    room: str = "Phòng Họp A"

class ChatMessage(BaseModel):
    message: str

class SuggestReq(BaseModel):
    problem: str

def suggest_roles_by_context(problem: str, existing_roles: List[Dict]) -> List[Dict]:
    # DE NHAN SU TU PHAN TICH NGU CANH + SAP XEP API THEO NGU CANH VAN DE
    # API manh ve khia canh nao thi giao vai do - KHONG CHIA DEU VONG TRON
    try:
        nhansu_provider = None
        for p in load_personnel():
            if p.get("core_type")=="nhansu":
                nhansu_provider = get_provider_info(p.get("provider","mistral"))
                nhansu_system = p.get("system_prompt","Bạn là Trưởng Phòng Nhân Sự. Đề xuất chức danh phù hợp và phân công API.")
                break
        if nhansu_provider is None:
            nhansu_provider = get_provider_info(ALIVE_PROVIDERS[0]["provider"] if ALIVE_PROVIDERS else "mistral")
            nhansu_system = "Bạn là Trưởng Phòng Nhân Sự - vai CORE quản lý Danh sách Hội đồng. Nhiệm vụ: phân tích ngữ cảnh vấn đề, đề xuất 2-3 chức danh phù hợp nhất (không cố định), và phân công API mạnh về khía cạnh đó cho từng vai. Chủ sự có quyền xoá và chỉnh sửa vai đó."

        existing_roles_str = ", ".join([r["role"] for r in existing_roles]) if existing_roles else "chưa có"
        providers_info = ", ".join([f"{p['provider']} (model {p['model']}) - mạnh về: {'lý luận, tổng hợp' if 'mistral' in p['provider'] else 'kỹ thuật, tốc độ, code' if 'groq' in p['provider'] else 'sáng tạo, marketing, ngôn ngữ'}" for p in ALIVE_PROVIDERS])
        
        prompt = f'''Vấn đề của Chủ sự: "{problem}"
Danh sách hội đồng hiện có: {existing_roles_str}
Các API đang có: {providers_info}

Nhiệm vụ của bạn (Nhân Sự - phụ trách thực hiện):
1. Phân tích ngữ cảnh vấn đề và đề xuất 2-4 CHỨC DANH phù hợp NHẤT cho vấn đề này. KHÔNG cố định phải là Kỹ Thuật + Marketing. Có thể là bất kỳ chuyên gia nào.
2. QUAN TRỌNG: Với mỗi chức danh, hãy phân công API mạnh về khía cạnh đó:
   - mistral: mạnh về lý luận, phân tích sâu, tổng hợp, văn hóa, tôn giáo
   - groq: mạnh về kỹ thuật, code, tốc độ, dữ liệu, tài chính, tính toán
   - openrouter: mạnh về sáng tạo, marketing, ngôn ngữ, giao tiếp

Ví dụ: Vấn đề về câu cá -> Chuyên gia Khí tượng giao cho groq (mạnh tính toán), Chuyên gia An toàn giao cho mistral (mạnh lý luận)
Vấn đề ngọc hoàng đại đế -> Chuyên gia Văn hóa Dân gian giao cho mistral (mạnh lý luận), Chuyên gia Truyền thông giao cho openrouter (mạnh sáng tạo)

Trả về đúng định dạng JSON array, mỗi phần tử có:
{{"role": "Tên chức danh (bắt đầu bằng Chuyên gia...)", "desc": "Mô tả ngắn 5-8 từ", "prompt": "System prompt chi tiết 50-80 từ", "provider": "mistral hoặc groq hoặc openrouter - API mạnh về khía cạnh đó"}}

Ví dụ: [{{"role": "Chuyên gia Văn hóa Dân gian", "desc": "Phân tích tín ngưỡng", "prompt": "Bạn là Chuyên gia Văn hóa...", "provider": "mistral"}}, {{"role": "Chuyên gia Marketing", "desc": "Chiến lược viral", "prompt": "...", "provider": "openrouter"}}]

Chỉ trả về JSON, không giải thích thêm.'''

        llm_response = call_llm_real(nhansu_provider, nhansu_system, prompt, max_tokens=800, temp=0.8)
        # Parse JSON tu response
        # Tim JSON array trong response
        import json as json_lib
        json_match = re.search(r'\[.*\]', llm_response, re.DOTALL)
        if json_match:
            try:
                parsed = json_lib.loads(json_match.group(0))
                if isinstance(parsed, list) and len(parsed)>0:
                    # Chuan hoa + Nhan Su sap xep API theo ngu canh
                    result = []
                    for item in parsed[:4]:
                        if isinstance(item, dict) and "role" in item:
                            role = item.get("role","").strip()
                            if not role.lower().startswith("chuyên gia"):
                                role = "Chuyên gia " + role
                            # Nhan Su phan cong API theo ngu canh van de
                            suggested_provider = item.get("provider","").lower()
                            if suggested_provider not in ["mistral","groq","openrouter","openai","gemini"]:
                                # Tu dong chon API manh ve khia canh do neu LLM khong tra provider
                                role_lower = role.lower()
                                if any(k in role_lower for k in ["kỹ thuật","code","dữ liệu","data","tài chính","tính toán","thủy văn","khí tượng"]):
                                    suggested_provider = "groq" if any(p["provider"]=="groq" for p in ALIVE_PROVIDERS) else ALIVE_PROVIDERS[0]["provider"] if ALIVE_PROVIDERS else "mistral"
                                elif any(k in role_lower for k in ["marketing","sáng tạo","truyền thông","ngôn ngữ","viral"]):
                                    suggested_provider = "openrouter" if any(p["provider"]=="openrouter" for p in ALIVE_PROVIDERS) else ALIVE_PROVIDERS[0]["provider"] if ALIVE_PROVIDERS else "mistral"
                                else:
                                    suggested_provider = "mistral" if any(p["provider"]=="mistral" for p in ALIVE_PROVIDERS) else ALIVE_PROVIDERS[0]["provider"] if ALIVE_PROVIDERS else "mistral"
                            result.append({
                                "role": role,
                                "desc": item.get("desc","Chuyên gia theo ngữ cảnh")[:80],
                                "prompt": item.get("prompt", f"Bạn là {role} theo chức danh. Đưa ý kiến chuyên môn chi tiết theo ngữ cảnh: {problem[:100]}")[:500],
                                "provider": suggested_provider,
                                "color": "#%06x" % (hash(role) % 0xFFFFFF)
                            })
                    if result:
                        print(f"🧠 Nhan Su de xuat dong (LLM): {[r['role'] for r in result]}")
                        return result
            except Exception as e:
                print(f"Parse JSON suggest failed: {e} - response: {llm_response[:200]}")
        
        # Fallback: dung keyword map neu LLM fail, nhung KHONG co dinh 2 nguoi
        print("Fallback keyword map cho suggest")
        pl=problem.lower(); suggested=[]; existing_names=[r["role"].lower() for r in existing_roles]
        km={
            "kỹ thuật|api|bug|code|lỗi|server|build|hệ thống|ai|sản phẩm|sagemaker|vertex|upsell": {"role": "Chuyên gia Kỹ Thuật", "desc": "Giải pháp kỹ thuật", "prompt": "Bạn là Chuyên gia Kỹ Thuật. Đưa giải pháp kỹ thuật cụ thể.", "color": "#6EC1E4"},
            "marketing|viral|khách hàng|tăng trưởng|quảng cáo|giá trị đơn hàng": {"role": "Chuyên gia Marketing", "desc": "Chiến lược viral", "prompt": "Bạn là Chuyên gia Marketing. Đưa chiến lược viral.", "color": "#FF8FA3"},
            "dữ liệu|phân tích|số liệu|báo cáo|metric|đánh giá khả thi|data": {"role": "Chuyên gia Dữ Liệu", "desc": "Phân tích số liệu", "prompt": "Bạn là Chuyên gia Dữ Liệu. Đưa số liệu, insight.", "color": "#A78BFA"},
            "tài chính|chi phí|tiền|roi|ngân sách|budget|tr": {"role": "Chuyên gia Tài Chính", "desc": "Tối ưu chi phí", "prompt": "Bạn là Chuyên gia Tài Chính. Phân tích chi phí, ROI.", "color": "#06D6A0"},
            "ngọc hoàng|đại đế|tôn giáo|tín ngưỡng|văn hóa dân gian|thần thoại": {"role": "Chuyên gia Văn hóa Dân gian & Tôn giáo", "desc": "Phân tích tín ngưỡng, thần thoại", "prompt": "Bạn là Chuyên gia Văn hóa Dân gian & Tôn giáo. Phân tích vấn đề dưới góc độ tín ngưỡng, thần thoại, văn hóa dân gian, lịch sử tôn giáo. Đưa góc nhìn học thuật, tôn trọng tín ngưỡng.", "color": "#F59E0B"},
            "khí tượng|thủy văn|câu cá|thời tiết|mưa giông|gió|nhiệt độ|thủy triều": {"role": "Chuyên gia Khí tượng Thủy văn & Hoạt động câu cá", "desc": "Đánh giá thời tiết, an toàn đi câu", "prompt": "Bạn là Chuyên gia Khí tượng Thủy văn & Hoạt động câu cá. Đánh giá mức độ phù hợp để đi câu dựa trên địa điểm, thời gian, dự báo mưa giông, tốc độ gió, nhiệt độ, thủy triều.", "color": "#06B6D4"},
            "pháp lý|luật|hợp đồng|tranh chấp": {"role": "Chuyên gia Pháp Lý", "desc": "Tư vấn pháp lý", "prompt": "Bạn là Chuyên gia Pháp Lý. Đưa tư vấn pháp lý cụ thể.", "color": "#8B5CF6"},
            "nhân sự|tuyển dụng|đào tạo": {"role": "Chuyên gia Nhân sự & Đào tạo", "desc": "Quản lý nhân sự", "prompt": "Bạn là Chuyên gia Nhân sự. Đưa giải pháp nhân sự.", "color": "#10B981"},
        }
        for pat, tpl in km.items():
            if re.search(pat, pl):
                if tpl["role"].lower() not in existing_names and tpl["role"] not in [s["role"] for s in suggested]:
                    suggested.append(tpl)
        if not suggested:
            # Neu khong khop keyword nao, de Nhan Su tu tao vai moi dua tren van de
            suggested = [{
                "role": f"Chuyên gia Phân tích Vấn đề",
                "desc": f"Chuyên gia cho: {problem[:30]}",
                "prompt": f"Bạn là Chuyên gia Phân tích Vấn đề theo ngữ cảnh. Vấn đề: {problem}. Đưa ý kiến chuyên môn phù hợp với ngữ cảnh này, không cố định.",
                "color": "#EC4899"
            }]
        return suggested[:3]
    except Exception as e:
        print(f"suggest_roles_by_context error: {e}")
        traceback.print_exc()
        return [{"role": "Chuyên gia Tổng hợp", "desc": "Phân tích tổng hợp", "prompt": f"Bạn là Chuyên gia Tổng hợp cho vấn đề: {problem[:100]}", "color": "#6EC1E4"}]

# ===== TIEN ICH TAO TEN BIEN BAN CHUAN VN =====
def make_bien_ban_filename(problem: str, ext: str = "md"):
    # Tom tat van de: lay 6-8 tu dau, loai bo ky tu dac biet
    # "biên bản họp - tóm tắt vấn đề - ngày tháng năm giờ phút giây"
    now = datetime.now()
    # Ngay thang nam gio phut giay: 02-09-2026 11-03-45
    time_str = now.strftime("%d-%m-%Y %H-%M-%S")
    # Tom tat: lay 40 ky tu dau, lam sach
    summary = problem.strip()[:50]
    # Loai bo ky tu dac biet giu lai chu, so, khoang trang
    summary = re.sub(r'[^\w\s\-]', '', summary, flags=re.UNICODE)
    summary = re.sub(r'\s+', ' ', summary).strip()
    if not summary:
        summary = "Cuoc hop hoi dong"
    # Ten chuan: bien ban hop - tom tat - ngay gio
    base_name = f"bien ban hop - {summary} - {time_str}"
    # File an toan (goc) van co the co dau tieng Viet, Windows cho phep
    # Nhung de an toan them ban khong dau
    safe_name = base_name
    # Thay / \ : * ? " < > | bang -
    safe_name = re.sub(r'[\/\:\*\?\"\<\>\|]', '-', safe_name)
    return f"{safe_name}.{ext}", base_name

def save_bien_ban_multi_format(minutes_content: str, problem: str):
    # Chi luu 2 dinh dang: txt, docx (bo md theo yeu cau Boss)
    base_txt, base_name = make_bien_ban_filename(problem, "txt")
    base_docx = base_txt.replace(".txt", ".docx")
    
    saved_files = []
    
    # 1. Luu .txt (cung noi dung)
    txt_path = DATA_DIR / base_txt
    txt_path.write_text(minutes_content, encoding='utf-8')
    saved_files.append(str(txt_path))
    
    # 2. Luu .docx (Word) neu co thu vien python-docx
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document()
        # Tieu de
        title = doc.add_heading(f'Biên bản họp - {problem[:60]}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Thong tin chung
        doc.add_paragraph(f'Thời gian: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        doc.add_paragraph(f'Vấn đề: {problem}')
        doc.add_paragraph(f'API: {len(ALIVE_PROVIDERS)} LLM ({", ".join([p["provider"] for p in ALIVE_PROVIDERS])})')
        doc.add_paragraph('---')
        
        # Noi dung - chia theo dong
        for line in minutes_content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), 1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), 2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), 3)
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                doc.add_paragraph(line.strip(), style='List Bullet')
            elif line.strip() == '---':
                doc.add_paragraph('_' * 50)
            else:
                if line.strip():
                    doc.add_paragraph(line)
        
        docx_path = DATA_DIR / base_docx
        doc.save(str(docx_path))
        saved_files.append(str(docx_path))
    except ImportError:
        print("python-docx chua cai, bo qua luu .docx - cai bang pip install python-docx")
    except Exception as e:
        print(f"Loi luu docx: {e}")
    
    return saved_files, base_name

@app.get("/")
def root():
    return {"name":"Council AI V20 FINAL V4 - Bien ban chuan VN + Word/TXT","alive_apis":len(ALIVE_PROVIDERS),"providers":[{"provider":p["provider"],"model":p["model"],"status":p.get("status")} for p in ALIVE_PROVIDERS],"personnel_count":len(load_personnel()),"note":"API=3 LLM, chuc danh=vo han, bien ban=ten chuan VN + 3 dinh dang md/txt/docx"}

@app.get("/personnel")
def get_personnel():
    return load_personnel()

@app.post("/personnel/suggest")
@app.post("/api/suggest")
@app.post("/suggest")
@app.post("/api/personnel/suggest")
def suggest_personnel(req: SuggestReq):
    existing=load_personnel()
    suggested=suggest_roles_by_context(req.problem, existing)
    # Tu dong them vao Danh sach Hoi dong luon theo yeu cau Boss
    auto_added = []
    current = load_personnel()
    existing_roles_lower = [p["role"].lower() for p in current]
    for idx, s in enumerate(suggested):
        if s["role"].lower() not in existing_roles_lower:
            nid = f"role_{uuid.uuid4().hex[:6]}"
            prov_idx = (len(current) + idx) % len(ALIVE_PROVIDERS) if ALIVE_PROVIDERS else 0
            prov = get_provider_info(ALIVE_PROVIDERS[prov_idx]["provider"] if ALIVE_PROVIDERS else "mistral")
            new_role = {
                "id": nid, "name": s["role"], "role": s["role"],
                "desc": s["desc"], "system_prompt": s["prompt"] + f" Ngữ cảnh: {req.problem[:300]}",
                "color": s["color"], "type": "council",
                "provider": prov["provider"], "model": prov["model"],
                "trained": True, "auto_generated": True,
                "created_at": datetime.now().isoformat()
            }
            current.append(new_role)
            auto_added.append(new_role)
    if auto_added:
        Path(PERSONNEL_FILE).write_text(json.dumps(current, ensure_ascii=False, indent=2), encoding='utf-8')
    return {"problem":req.problem,"suggested":suggested,"auto_added":auto_added,"auto_call":True, "message": f"Đã tự thêm {len(auto_added)} vai vào Danh sách Hội đồng - Chủ sự có quyền xoá và chỉnh sửa"}

@app.post("/personnel/suggest-legacy")
def suggest_legacy(req: dict):
    problem=req.get("problem","")
    existing=load_personnel()
    suggested=suggest_roles_by_context(problem, existing)
    # Auto add
    auto_added = []
    current = load_personnel()
    existing_roles_lower = [p["role"].lower() for p in current]
    for idx, s in enumerate(suggested):
        if s["role"].lower() not in existing_roles_lower:
            nid = f"role_{uuid.uuid4().hex[:6]}"
            prov_idx = (len(current) + idx) % len(ALIVE_PROVIDERS) if ALIVE_PROVIDERS else 0
            prov = get_provider_info(ALIVE_PROVIDERS[prov_idx]["provider"] if ALIVE_PROVIDERS else "mistral")
            new_role = {
                "id": nid, "name": s["role"], "role": s["role"],
                "desc": s["desc"], "system_prompt": s["prompt"] + f" Ngữ cảnh: {problem[:300]}",
                "color": s["color"], "type": "council",
                "provider": prov["provider"], "model": prov["model"],
                "trained": True, "auto_generated": True,
                "created_at": datetime.now().isoformat()
            }
            current.append(new_role)
            auto_added.append(new_role)
    if auto_added:
        Path(PERSONNEL_FILE).write_text(json.dumps(current, ensure_ascii=False, indent=2), encoding='utf-8')
    return {"problem":problem,"suggested":suggested,"auto_added":auto_added,"auto_call":True}

@app.get("/app")
@app.get("/frontend")
@app.get("/ui")
def serve_frontend():
    if FRONTEND_FILE.exists():
        return FileResponse(FRONTEND_FILE, media_type="text/html")
    return {"error":"Chua co council-v20-frontend.html"}

councils_db: Dict = {}

@app.post("/councils")
def create_council(req: CreateCouncilReq):
    try:
        cid=uuid.uuid4().hex[:8]
        all_personnel={p["id"]:p for p in load_personnel()}
        core_ids=["core_thuky","core_phoql","core_nhansu"]
        selected=set(req.selected_personnel_ids)
        for c in core_ids: selected.add(c)
        suggested=suggest_roles_by_context(req.problem, [all_personnel[i] for i in selected if i in all_personnel])
        auto_added=[]
        for s in suggested:
            exists=any(all_personnel[pid]["role"].lower()==s["role"].lower() for pid in selected if pid in all_personnel)
            if not exists:
                nid=f"role_{uuid.uuid4().hex[:6]}"
                prov_idx = (len(all_personnel) + len(auto_added)) % len(ALIVE_PROVIDERS) if ALIVE_PROVIDERS else 0
                prov=get_provider_info(ALIVE_PROVIDERS[prov_idx]["provider"] if ALIVE_PROVIDERS else "mistral")
                new_role={"id":nid,"name":s["role"],"role":s["role"],"desc":s["desc"],"system_prompt":s["prompt"]+f" Ngữ cảnh: {req.problem[:300]}","color":s["color"],"type":"council","provider":prov["provider"],"model":prov["model"],"trained":True,"auto_generated":True,"created_at":datetime.now().isoformat()}
                cur=load_personnel(); cur.append(new_role); Path(PERSONNEL_FILE).write_text(json.dumps(cur, ensure_ascii=False, indent=2), encoding='utf-8')
                all_personnel[nid]=new_role; selected.add(nid); auto_added.append(new_role)
        roles={}
        for pid in selected:
            if pid not in all_personnel: continue
            p=all_personnel[pid]
            roles[pid]={"short_name":p["name"],"role":p["role"],"system":p["system_prompt"],"color":p["color"],"provider_info":get_provider_info(p["provider"]),"trained":p.get("trained",False),"core_type":p.get("core_type",""),"is_core":p.get("type")=="core"}
        councils_db[cid]={"id":cid,"title":req.problem[:80],"problem":req.problem,"room":req.room,"roles":roles,"active_roles":list(roles.keys()),"history":[],"minutes":[],"created_at":datetime.now().isoformat(),"auto_added":auto_added}
        return {"id":cid,"title":req.problem[:80],"roles":list(roles.values()),"active_roles":list(roles.keys()),"auto_added":auto_added}
    except Exception as e:
        return {"error":str(e), "detail": traceback.format_exc()[:1000]}

@app.post("/councils/{cid}/run")
async def run_council(cid: str):
    try:
        if cid not in councils_db: return {"error":"Council not found"}
        data=councils_db[cid]
        thuky=next((v for k,v in data["roles"].items() if v.get("core_type")=="thuky"), None)
        rollcall = await asyncio.to_thread(call_llm_real, thuky["provider_info"] if thuky else get_provider_info("mistral"), thuky["system"] if thuky else "Ban la Thu Ky", f"Vấn đề: '{data['problem']}' - Điểm danh {len(data['roles'])} vai: {', '.join([r['short_name'] for r in data['roles'].values()])}", 500, 0.6) if thuky else f"Điểm danh: {', '.join([r['short_name'] for r in data['roles'].values()])}"

        experts=[]
        for rid,info in data["roles"].items():
            if info.get("core_type") in ["thuky","phoql","nhansu"]: continue
            if rid not in data["active_roles"]: continue
            try:
                ans=await asyncio.to_thread(call_llm_real, info["provider_info"], info["system"], f"Vấn đề: '{data['problem']}' - Bạn là {info['short_name']} ({info['role']}). Đưa ý kiến đầy đủ chi tiết 250-400 từ, không cắt cụt.", 600, 0.8)
                experts.append({"personnel_id":rid,"name":info["short_name"],"role":info["role"],"content":ans,"provider":info["provider_info"]["provider"]})
            except Exception as e:
                experts.append({"personnel_id":rid,"name":info["short_name"],"role":info["role"],"content":f"Fallback {info['short_name']}: {str(e)[:80]}","provider":info["provider_info"]["provider"]})

        nhansu=next((v for k,v in data["roles"].items() if v.get("core_type")=="nhansu"), None)
        hr_suggest = await asyncio.to_thread(call_llm_real, nhansu["provider_info"], nhansu["system"], f"Vấn đề: '{data['problem']}' - Hội đồng: {', '.join([r['short_name'] for r in data['roles'].values()])} - Thiếu chức danh nào?", 400, 0.7) if nhansu else "Đề xuất bổ sung Vận Hành"

        phoql=next((v for k,v in data["roles"].items() if v.get("core_type")=="phoql"), None)
        expert_full="\n".join([f"- {e['name']} ({e['role']}): {e['content']}" for e in experts])
        summary = await asyncio.to_thread(call_llm_real, phoql["provider_info"], phoql["system"], f"Vấn đề: '{data['problem']}'\nBiên bản: {rollcall}\nÝ kiến đầy đủ:\n{expert_full}\nNhân Sự: {hr_suggest}\nRa quyết định cuối cùng đầy đủ chi tiết, 3 bước hành động có deadline, phân công theo chức danh. Đây chính là biên bản chính.", 800, 0.75) if phoql else "Quyết định fallback"

        minutes_content = f"# BIÊN BẢN HỌP - {data['title']}\n\n**Thời gian:** {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n**Phòng:** {data['room']}\n**Vấn đề:** {data['problem']}\n**API:** {len(ALIVE_PROVIDERS)} LLM ({', '.join([p['provider'] for p in ALIVE_PROVIDERS])}) - Chức danh: {len(data['roles'])} vai (vô hạn, do Nhân Sự đặt)\n\n## 1. Điểm danh (Thư Ký)\n{rollcall}\n\n## 2. Ý kiến Hội Đồng (Trọn vẹn)\n"
        for e in experts:
            minutes_content += f"\n### {e['name']} - {e['role']} [{e['provider']}]\n{e['content']}\n\n"
        minutes_content += f"## 3. Nhân Sự đề xuất\n{hr_suggest}\n\n## 4. Quyết định cuối cùng - Phó Quản Lý Điều Hành (Đây là biên bản chính)\n{summary}\n\n---\n*Tổng {len(minutes_content)} ký tự - Lưu trọn vẹn*\n"

        # Luu voi ten chuan VN + 3 dinh dang
        saved_files, base_name = save_bien_ban_multi_format(minutes_content, data["problem"])
        print(f"✅ Da luu bien ban: {base_name} -> {saved_files}")

        result={"id":cid,"problem":data["problem"],"rollcall":rollcall,"experts":experts,"hr_suggest":hr_suggest,"summary":summary,"minutes_file":saved_files[0] if saved_files else "","minutes_files":saved_files,"minutes_content":minutes_content,"base_name":base_name,"roles_used":[r["short_name"] for r in data["roles"].values()],"auto_added":data.get("auto_added",[]),"total_chars":len(minutes_content),"is_full":True,"api_count":len(ALIVE_PROVIDERS),"chuc_danh_count":len(data["roles"])}
        councils_db[cid]["last_result"]=result
        return result
    except Exception as e:
        return {"error":str(e), "detail": traceback.format_exc()[:2000]}

@app.post("/councils/{cid}/chat")
async def continue_chat(cid: str, msg: ChatMessage):
    try:
        if cid not in councils_db: return {"error":"Not found"}
        data=councils_db[cid]
        clean=msg.message.replace("@@","@")
        clean_lower=clean.lower()
        
        # Neu @Thư Ký và co ten chuyen gia khong co trong cuoc hop -> goi vao theo V4
        # Vi du: @Thư Ký Hội Đồng gọi Chuyên gia Khí tượng vào
        if "@thư ký" in clean_lower or "@thu ky" in clean_lower:
            # Tim tat ca chuc danh trong personnel co trong tin nhan
            all_personnel = {p["id"]: p for p in load_personnel()}
            added = []
            for pid, p in all_personnel.items():
                role_lower = p["role"].lower()
                # Neu ten role xuat hien trong message va chua co trong active
                if role_lower in clean_lower and pid not in data["active_roles"]:
                    # Them vao cuoc hop
                    if pid not in data["roles"]:
                        # Them vao roles dict
                        data["roles"][pid] = {
                            "short_name": p["name"], "role": p["role"],
                            "system": p["system_prompt"], "color": p["color"],
                            "provider_info": get_provider_info(p["provider"]),
                            "trained": p.get("trained",False),
                            "core_type": p.get("core_type",""), "is_core": p.get("type")=="core"
                        }
                    data["active_roles"].append(pid)
                    added.append(p["role"])
            if added:
                return {"action":"add_team","added_roles":added,"kept_roles":[data["roles"][rid]["short_name"] for rid in data["active_roles"]],"active_roles":data["active_roles"],"type":"text","name":"Thư Ký Hội Đồng","short_name":"Thư Ký Hội Đồng","color":"#2B2D42","text":f"Đã gọi thêm vào cuộc họp theo yêu cầu @Thư Ký: {', '.join(added)}. Team hiện tại: {', '.join([data['roles'][rid]['short_name'] for rid in data['active_roles']])}. (Nhân sự phụ trách xử lý theo trình bày giới thiệu bản V4)"}
        
        # @ tag hien du vai trong cuoc hop
        # Tim mention trong message
        mentioned_ids = []
        for rid, info in data["roles"].items():
            if f"@{info['short_name']}" in clean or f"@{info['role']}" in clean:
                if rid in data["active_roles"]:
                    mentioned_ids.append(rid)
        
        if mentioned_ids:
            results = {}
            for rid in mentioned_ids[:6]:
                info = data["roles"][rid]
                ans = await asyncio.to_thread(call_llm_real, info["provider_info"], info["system"], f"Chủ sự hỏi riêng bạn: '{clean}' - Vấn đề gốc: '{data['problem']}' - Trả lời 150 từ, theo chức danh {info['role']}.", 300, 0.85)
                results[rid] = {"type":"text","text":ans,"name":info["short_name"],"short_name":info["short_name"],"color":info["color"],"provider":info["provider_info"]["provider"]}
            return {"action":"mention","mentioned":True,"results":results}
        
        summary_role=next((v for v in data["roles"].values() if v.get("core_type")=="phoql"), None) or next(iter(data["roles"].values()))
        ans=await asyncio.to_thread(call_llm_real, summary_role["provider_info"], summary_role["system"], f"Chủ sự hỏi: '{clean}' - Vấn đề gốc: '{data['problem']}' - Trả lời 150 từ.", 300, 0.8)
        return {"action":"general","type":"text","name":summary_role["short_name"],"short_name":summary_role["short_name"],"color":summary_role["color"],"text":ans}
    except Exception as e:
        traceback.print_exc()
        return {"action":"general","type":"text","name":"Phó Quản Lý","short_name":"Phó Quản Lý","color":"#EF476F","text":f"Fallback: {str(e)[:100]}"}

@app.get("/councils/{cid}/minutes")
def get_minutes(cid: str):
    if cid not in councils_db: return {"error":"Not found"}
    data=councils_db[cid]
    res=data.get("last_result",{})
    return {"council_id":cid,"title":data["title"],"problem":data["problem"],"minutes_file":res.get("minutes_file"),"minutes_files":res.get("minutes_files",[]),"minutes_content":res.get("minutes_content",""),"total_chars":res.get("total_chars",0),"is_full":True,"summary":res.get("summary",""),"base_name":res.get("base_name","")}

@app.get("/meetings")
def list_meetings():
    files = list(MEETINGS_DIR.glob("*.*"))
    meetings = []
    for f in sorted(files, key=lambda x: x.stat().st_mtime, reverse=True)[:30]:
        meetings.append({"file": str(f), "name": f.name, "size": f.stat().st_size, "modified": datetime.fromtimestamp(f.stat().st_mtime).isoformat(), "ext": f.suffix})
    return {"meetings": meetings, "count": len(files)}

@app.get("/admin/fix-personnel")
def fix_personnel_get():
    # Fix for Boss when personnel.json corrupted
    if PERSONNEL_FILE.exists():
        try:
            content = PERSONNEL_FILE.read_text(encoding='utf-8').strip()
            if not content:
                raise ValueError("empty")
            json.loads(content)
            return {"ok": True, "message": "personnel.json OK"}
        except:
            PERSONNEL_FILE.unlink(missing_ok=True)
    data=get_default_personnel()
    Path(PERSONNEL_FILE).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
    return {"ok": True, "message": "Đã fix personnel.json bị lỗi Expecting value", "data": data}

@app.post("/admin/clean-personnel")
def clean_personnel():
    if PERSONNEL_FILE.exists(): PERSONNEL_FILE.unlink()
    data=get_default_personnel()
    Path(PERSONNEL_FILE).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
    return {"ok":True,"data":data}

@app.post("/admin/reload-apis")
def reload_apis():
    refresh_alive_providers()
    return {"ok":True,"alive":len(ALIVE_PROVIDERS),"providers":ALIVE_PROVIDERS}

if __name__=="__main__":
    import uvicorn
    print("\n" + "="*60)
    print("🏢 COUNCIL AI V20 FINAL V4 - Bien ban chuan VN + Word/TXT")
    print("   Ten: bien ban hop - tom tat van de - dd-mm-yyyy hh-mm-ss.ext")
    print("   Chi luu 2 dinh dang: .txt, .docx (neu co python-docx)")
    print("="*60)
    uvicorn.run(app, host="0.0.0.0", port=8000)
